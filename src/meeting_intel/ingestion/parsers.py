import re
from pathlib import Path

from meeting_intel.schemas import MeetingDocument, SourceType, TranscriptTurn


SPEAKER_LINE = re.compile(r"^\s*(?P<speaker>[A-Za-z][\w .'-]{0,60})\s*:\s*(?P<text>.+)$")


def parse_transcript_text(
    text: str,
    title: str = "Untitled Meeting",
    source_type: SourceType = SourceType.raw_text,
    participants: list[str] | None = None,
) -> MeetingDocument:
    turns: list[TranscriptTurn] = []
    inferred_participants = set(participants or [])

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = SPEAKER_LINE.match(line)
        if match:
            speaker = normalize_speaker(match.group("speaker"))
            content = clean_text(match.group("text"))
            inferred_participants.add(speaker)
            turns.append(TranscriptTurn(speaker=speaker, text=content))
        elif turns:
            turns[-1].text = clean_text(f"{turns[-1].text} {line}")
        else:
            turns.append(TranscriptTurn(text=clean_text(line)))

    return MeetingDocument(
        title=title,
        participants=sorted(inferred_participants),
        transcript=turns,
        source_type=source_type,
    )


def normalize_meetingbank_record(record: dict) -> MeetingDocument:
    meeting_id = str(record.get("id") or record.get("meeting_id") or "")
    title = str(record.get("title") or record.get("source") or f"MeetingBank {meeting_id}").strip()
    transcript = str(
        record.get("transcript")
        or record.get("meeting_transcript")
        or record.get("text")
        or record.get("dialogue")
        or ""
    )
    summary = str(record.get("summary") or record.get("reference_summary") or "").strip()
    document = parse_transcript_text(transcript, title=title, source_type=SourceType.meetingbank)
    document.summary = summary
    document.embeddings_metadata["source_dataset"] = "meetingbank"
    if meeting_id:
        document.embeddings_metadata["source_record_id"] = meeting_id
    return document


def parse_file(path: Path, title: str | None = None) -> MeetingDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path, title or path.stem)
    if suffix == ".docx":
        return parse_docx(path, title or path.stem)
    return parse_transcript_text(path.read_text(encoding="utf-8"), title or path.stem, SourceType.txt)


def parse_pdf(path: Path, title: str) -> MeetingDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return parse_transcript_text(text, title, SourceType.pdf)


def parse_docx(path: Path, title: str) -> MeetingDocument:
    from docx import Document

    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return parse_transcript_text(text, title, SourceType.docx)


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def normalize_speaker(speaker: str) -> str:
    speaker = clean_text(speaker)
    return speaker.title() if speaker.isupper() else speaker
